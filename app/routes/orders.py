import os
import re
import json
from datetime import datetime, timezone
from urllib.parse import urlparse

from flask import Blueprint, abort, current_app, flash, jsonify, redirect, render_template, request, send_file, url_for
from flask_login import current_user, login_required
from werkzeug.utils import secure_filename

from ..extensions import db
from ..models import OrderDocument
from ..services.ai_client import ask_openai, is_ai_unavailable_message
from ..services.orders_ingestion import get_ingestion_state, run_official_orders_ingestion


bp = Blueprint('orders', __name__)

ORDER_UPLOAD_DIR = os.path.join('data', 'uploads', 'orders')
ORDER_SEED_PATH = os.path.join('data', 'orders', 'seed_orders.json')
APPROVED_SOURCES_PATH = os.path.join('data', 'orders', 'approved_sources.json')
ALLOWED_EXTENSIONS = {'.pdf', '.doc', '.docx', '.txt', '.rtf'}
SOURCE_TYPE_OPTIONS = [
    'MCO',
    'MCBUL',
    'MARADMIN',
    'ALMAR',
    'NAVMC',
    'DOD_DIRECTIVE',
    'MCLB_ALBANY_ORDER',
    'USMC_ORDER',
    'MEMORANDUM',
    'LOCAL_DOCUMENT',
]
ALLOWED_SEARCH_SOURCE_TYPES = set(SOURCE_TYPE_OPTIONS) | {'USMC_ORDER'}
DEFAULT_APPROVED_DOMAINS = (
    'marines.mil',
    'usmc.mil',
    'navy.mil',
    'esd.whs.mil',
    'defense.gov',
    'dod.mil',
    'govinfo.gov',
    'uscode.house.gov',
    'ecfr.gov',
)
DEFAULT_APPROVED_SOURCE_GROUPS = (
    'USMC',
    'USMC Official',
    'MCLB Albany',
    'PMO',
    'Security',
    'HR',
    'Internal Upload',
    'Seed Library',
)
BLOCKED_SOURCE_HINTS = (
    'indeed',
    'linkedin',
    'job',
    'jobs',
    'career',
    'careers',
    'blog',
    'forum',
    'reddit',
    'facebook',
    'instagram',
    'x.com',
    'twitter',
    'youtube',
    'shop',
    'amazon',
)


def _empty_order_ai_strategy():
    return {
        'priority_terms': [],
        'query_variants': [],
        'topic_hints': [],
        'audience_hints': [],
        'officer_brief': '',
    }


def _dedupe_search_phrases(values, limit=8):
    cleaned = []
    seen = set()
    for value in values or []:
        text = re.sub(r'\s+', ' ', str(value or '').strip())
        if not text:
            continue
        key = text.lower()
        if key in seen:
            continue
        seen.add(key)
        cleaned.append(text)
        if len(cleaned) >= limit:
            break
    return cleaned


def _extract_json_object(text):
    raw = (text or '').strip()
    if not raw:
        return {}
    fenced = re.search(r'```(?:json)?\s*(\{.*?\})\s*```', raw, flags=re.S)
    if fenced:
        raw = fenced.group(1).strip()
    if not raw.startswith('{'):
        start = raw.find('{')
        end = raw.rfind('}')
        if start >= 0 and end > start:
            raw = raw[start:end + 1]
    try:
        parsed = json.loads(raw)
        return parsed if isinstance(parsed, dict) else {}
    except json.JSONDecodeError:
        return {}


def _normalize_ai_strategy(payload):
    strategy = _empty_order_ai_strategy()
    if not isinstance(payload, dict):
        return strategy
    strategy['priority_terms'] = _dedupe_search_phrases(payload.get('priority_terms') or [], limit=8)
    strategy['query_variants'] = _dedupe_search_phrases(payload.get('query_variants') or [], limit=4)
    strategy['topic_hints'] = _dedupe_search_phrases(payload.get('topic_hints') or [], limit=6)
    strategy['audience_hints'] = _dedupe_search_phrases(payload.get('audience_hints') or [], limit=4)
    strategy['officer_brief'] = str(payload.get('officer_brief') or '').strip()
    return strategy

try:
    from pypdf import PdfReader
except Exception:  # pragma: no cover
    PdfReader = None

try:
    from docx import Document
except Exception:  # pragma: no cover
    Document = None


def _orders_directory():
    root = current_app.root_path
    directory = os.path.join(root, ORDER_UPLOAD_DIR)
    os.makedirs(directory, exist_ok=True)
    return directory


def _safe_order_file_path(path):
    raw_path = (path or '').strip()
    if not raw_path:
        return ''
    candidate = os.path.abspath(raw_path)
    app_root = os.path.abspath(current_app.root_path)
    try:
        if os.path.commonpath([candidate, app_root]) != app_root:
            return ''
    except ValueError:
        return ''
    if not os.path.isfile(candidate):
        return ''
    return candidate


def _document_download_path(document):
    return _safe_order_file_path(document.file_path)


def _document_has_download(document):
    return bool(_document_download_path(document))


def _order_seed_file_path():
    root = current_app.root_path
    return os.path.join(root, ORDER_SEED_PATH)


def _approved_sources_file_path():
    root = current_app.root_path
    return os.path.join(root, APPROVED_SOURCES_PATH)


def _default_approved_sources_config():
    return {
        'approved_domains': list(DEFAULT_APPROVED_DOMAINS),
        'approved_source_groups': list(DEFAULT_APPROVED_SOURCE_GROUPS),
        'blocked_hints': list(BLOCKED_SOURCE_HINTS),
    }


def _normalize_text_list(values):
    normalized = []
    for value in values or []:
        cleaned = str(value or '').strip()
        if cleaned and cleaned not in normalized:
            normalized.append(cleaned)
    return normalized


def _load_approved_sources_config():
    path = _approved_sources_file_path()
    defaults = _default_approved_sources_config()
    if not os.path.exists(path):
        return defaults
    try:
        with open(path, 'r', encoding='utf-8') as handle:
            payload = json.load(handle)
    except (OSError, json.JSONDecodeError):
        return defaults
    if not isinstance(payload, dict):
        return defaults
    return {
        'approved_domains': _normalize_text_list(payload.get('approved_domains') or defaults['approved_domains']),
        'approved_source_groups': _normalize_text_list(payload.get('approved_source_groups') or defaults['approved_source_groups']),
        'blocked_hints': _normalize_text_list(payload.get('blocked_hints') or defaults['blocked_hints']),
    }


def _save_approved_sources_config(config):
    path = _approved_sources_file_path()
    os.makedirs(os.path.dirname(path), exist_ok=True)
    payload = {
        'approved_domains': sorted(_normalize_text_list(config.get('approved_domains') or []), key=str.lower),
        'approved_source_groups': sorted(_normalize_text_list(config.get('approved_source_groups') or []), key=str.lower),
        'blocked_hints': sorted(_normalize_text_list(config.get('blocked_hints') or []), key=str.lower),
    }
    with open(path, 'w', encoding='utf-8') as handle:
        json.dump(payload, handle, indent=2)


def _extract_source_urls(value):
    return re.findall(r'https?://[^\s|,]+', value or '')


def _normalized_host(value):
    parsed = urlparse(value if '://' in (value or '') else f'https://{value}')
    return (parsed.netloc or parsed.path or '').strip().lower().lstrip('.')


def _is_approved_domain(host, config):
    candidate = _normalized_host(host)
    if not candidate:
        return False
    for blocked in config.get('blocked_hints') or []:
        if blocked.lower() in candidate:
            return False
    for approved in config.get('approved_domains') or []:
        normalized = _normalized_host(approved)
        if candidate == normalized or candidate.endswith(f'.{normalized}'):
            return True
    return False


def _is_local_orders_file(path):
    raw_path = (path or '').strip()
    if not raw_path:
        return False
    candidate = os.path.abspath(raw_path)
    root = os.path.abspath(_orders_directory())
    try:
        return os.path.commonpath([candidate, root]) == root
    except ValueError:
        return False


def _source_group_allowed(source_group, config):
    raw = (source_group or '').strip()
    if not raw:
        return False
    lowered = raw.lower()
    if any(blocked.lower() in lowered for blocked in (config.get('blocked_hints') or [])):
        return False
    for approved in config.get('approved_source_groups') or []:
        approved_clean = approved.strip()
        if not approved_clean:
            continue
        if raw == approved_clean or lowered.startswith(f'{approved_clean.lower()} |'):
            return True
    return False


def _document_source_metadata(document, config=None):
    config = config or _load_approved_sources_config()
    source_group = (document.source_group or '').strip()
    urls = _extract_source_urls(source_group)
    hosts = []
    for url in urls:
        host = _normalized_host(url)
        if host and host not in hosts:
            hosts.append(host)
    file_name = os.path.basename((document.file_path or '').strip()).lower()
    is_local = _is_local_orders_file(document.file_path)
    is_seed = file_name.startswith('seed-') or (document.source_version or '').lower().startswith('seed')
    source_type = (document.source_type or '').strip().upper()
    blocked_reason = ''
    combined_source_text = ' '.join([source_group, ' '.join(hosts)]).lower()
    for blocked in config.get('blocked_hints') or []:
        if blocked.lower() in combined_source_text:
            blocked_reason = f'blocked hint: {blocked}'
            break

    approved = False
    origin_type = 'unknown'
    origin_label = source_group or 'Unknown'

    if blocked_reason:
        approved = False
    elif source_type and source_type not in ALLOWED_SEARCH_SOURCE_TYPES:
        blocked_reason = f'unapproved source type: {source_type}'
        approved = False
    elif hosts:
        approved = all(_is_approved_domain(host, config) for host in hosts)
        origin_type = 'domain'
        origin_label = hosts[0]
        if not approved:
            blocked_reason = 'domain not allowlisted'
    elif is_seed:
        approved = True
        origin_type = 'seed'
        origin_label = 'Seed Library'
    elif _source_group_allowed(source_group, config):
        approved = True
        origin_type = 'source_group'
        origin_label = source_group
    elif is_local and (document.uploaded_by is not None or not source_group):
        approved = True
        origin_type = 'local'
        origin_label = source_group or 'Internal Upload'
    elif is_local and _source_group_allowed(source_group, config):
        approved = True
        origin_type = 'local'
        origin_label = source_group
    else:
        blocked_reason = blocked_reason or 'source is not on approved allowlist'

    return {
        'approved': approved,
        'origin_type': origin_type,
        'origin_label': origin_label,
        'hosts': hosts,
        'is_local': is_local,
        'is_seed': is_seed,
        'blocked_reason': blocked_reason,
    }


def _approved_documents():
    config = _load_approved_sources_config()
    documents = OrderDocument.query.order_by(OrderDocument.uploaded_at.desc(), OrderDocument.id.desc()).all()
    approved = []
    for document in documents:
        metadata = _document_source_metadata(document, config)
        if metadata['approved']:
            document.source_origin_label = metadata['origin_label']
            document.source_origin_type = metadata['origin_type']
            approved.append(document)
    return approved


def _write_seed_text_file(slug, text_content):
    safe_slug = re.sub(r'[^a-z0-9\-]+', '-', slug.lower()).strip('-') or f"order-{int(datetime.utcnow().timestamp())}"
    filename = f"seed-{safe_slug}.txt"
    path = os.path.join(_orders_directory(), filename)
    if not os.path.exists(path):
        with open(path, 'w', encoding='utf-8') as handle:
            handle.write(text_content.strip() + '\n')
    return path


def _load_seed_orders():
    seed_path = _order_seed_file_path()
    if not os.path.exists(seed_path):
        return []
    with open(seed_path, 'r', encoding='utf-8') as handle:
        payload = json.load(handle)
    if isinstance(payload, dict):
        return payload.get('orders', []) or []
    if isinstance(payload, list):
        return payload
    return []


def _seed_orders_library():
    seed_rows = _load_seed_orders()
    if not seed_rows:
        return 0, 0

    inserted = 0
    skipped = 0
    for row in seed_rows:
        title = (row.get('title') or '').strip()
        if not title:
            skipped += 1
            continue

        source_type = (row.get('source_type') or 'MEMORANDUM').strip().upper()
        if source_type not in SOURCE_TYPE_OPTIONS:
            source_type = 'MEMORANDUM'

        order_number = (row.get('order_number') or '').strip() or None
        memo_number = (row.get('memo_number') or '').strip() or None
        existing = (
            OrderDocument.query
            .filter_by(title=title, source_type=source_type, order_number=order_number, memo_number=memo_number)
            .first()
        )
        if existing:
            skipped += 1
            continue

        aliases = row.get('aliases') or []
        aliases_text = ', '.join(str(item).strip() for item in aliases if str(item).strip())
        audience_tags = ', '.join(str(item).strip() for item in (row.get('audience_tags') or []) if str(item).strip())
        topic_tags = ', '.join(str(item).strip() for item in (row.get('topic_tags') or []) if str(item).strip())
        summary = (row.get('summary') or '').strip() or None
        category = (row.get('category') or '').strip() or None
        issuing_authority = (row.get('issuing_authority') or '').strip() or None
        source_group = (row.get('source_group') or '').strip() or None
        source_version = (row.get('source_version') or 'seed-v1').strip()
        issue_date = _parse_date(row.get('issue_date'))
        revision_date = _parse_date(row.get('revision_date'))

        full_text = (row.get('full_text') or summary or title).strip()
        extracted_text_parts = [full_text]
        if aliases_text:
            extracted_text_parts.append(f"Aliases: {aliases_text}")
        if audience_tags:
            extracted_text_parts.append(f"Audience: {audience_tags}")
        if topic_tags:
            extracted_text_parts.append(f"Topics: {topic_tags}")
        extracted_text_parts.extend(
            [
                f"Title: {title}",
                f"Category: {category or 'General'}",
                f"Order Number: {order_number or '-'}",
                f"Memo Number: {memo_number or '-'}",
                f"Issuing Authority: {issuing_authority or '-'}",
                f"Summary: {summary or '-'}",
            ]
        )
        extracted_text = '\n'.join(part for part in extracted_text_parts if part)

        slug = (order_number or memo_number or title)[:80]
        file_path = _write_seed_text_file(slug, extracted_text)

        db.session.add(
            OrderDocument(
                title=title,
                category=category,
                source_type=source_type,
                source_group=source_group,
                order_number=order_number,
                memo_number=memo_number,
                issuing_authority=issuing_authority,
                issue_date=issue_date,
                revision_date=revision_date,
                source_version=source_version,
                audience_tags=audience_tags or None,
                topic_tags=topic_tags or None,
                version_label=(row.get('version_label') or '').strip() or source_version,
                summary=summary,
                extracted_text=extracted_text,
                parser_confidence=0.95,
                file_path=file_path,
                uploaded_by=current_user.id if getattr(current_user, 'is_authenticated', False) else None,
                last_indexed_at=datetime.now(timezone.utc).replace(tzinfo=None),
                is_active=True,
            )
        )
        inserted += 1

    if inserted:
        db.session.commit()
    return inserted, skipped


def _ensure_orders_library_seeded():
    if OrderDocument.query.count() > 0:
        return 0, 0
    if not _load_seed_orders():
        return 0, 0
    return _seed_orders_library()


def _safe_display_dt(value):
    if not value:
        return ''
    return value.strftime('%Y-%m-%d %H:%M:%S ET')


def _parse_date(value):
    raw = (value or '').strip()
    if not raw:
        return None
    for fmt in ('%Y-%m-%d', '%m/%d/%Y'):
        try:
            return datetime.strptime(raw, fmt)
        except ValueError:
            continue
    return None


def _extract_text_from_file(path):
    extension = os.path.splitext(path)[1].lower()
    try:
        if extension in {'.txt', '.rtf'}:
            with open(path, 'r', encoding='utf-8', errors='ignore') as handle:
                return handle.read()
        if extension == '.pdf' and PdfReader:
            reader = PdfReader(path)
            chunks = []
            for page in reader.pages[:20]:
                chunks.append(page.extract_text() or '')
            return '\n'.join(chunks).strip()
        if extension == '.docx' and Document:
            doc = Document(path)
            return '\n'.join(paragraph.text for paragraph in doc.paragraphs if paragraph.text).strip()
    except Exception:
        return ''
    return ''


def _tokenize_search_text(value):
    cleaned = re.sub(r'[^a-z0-9\s]', ' ', (value or '').lower())
    return [token for token in cleaned.split() if len(token) > 2]


_ORDER_SYNONYM_MAP = {
    'dui': ['alcohol', 'intoxicated', 'driving', 'vehicle'],
    'drunk': ['alcohol', 'intoxicated'],
    'gate': ['access', 'entry', 'visitor', 'search'],
    'weapon': ['firearm', 'armory', 'rifle', 'pistol'],
    'weapons': ['firearm', 'armory', 'rifle', 'pistol'],
    'domestic': ['family', 'spouse', 'violence', 'protective', 'order'],
    'fight': ['assault', 'battery', 'disturbance'],
    'barracks': ['housing', 'room', 'quarters'],
    'vehicle': ['traffic', 'driving', 'registration', 'inspection'],
    'memo': ['memorandum', 'policy', 'directive'],
    'memorandum': ['memo', 'policy', 'directive'],
    'orders': ['policy', 'directive', 'instruction'],
    'policy': ['directive', 'instruction', 'guidance'],
    'guidance': ['policy', 'instruction', 'memorandum'],
    'hr': ['personnel', 'civilian', 'leave', 'attendance'],
    'personnel': ['hr', 'leave', 'attendance', 'civilian'],
    'leave': ['absence', 'vacation', 'personnel'],
    'civilian': ['employee', 'personnel', 'administrative'],
    'employee': ['civilian', 'personnel', 'administrative'],
    'pay': ['finance', 'personnel', 'administrative'],
    'it': ['network', 'cyber', 'computer', 'access'],
    'computer': ['it', 'network', 'cyber'],
    'maintenance': ['facility', 'workorder', 'safety'],
    'facility': ['maintenance', 'workorder', 'safety'],
    'housing': ['barracks', 'quarters', 'occupancy'],
    'travel': ['tdy', 'orders', 'transportation'],
    'training': ['qualification', 'certification', 'instruction'],
    'safety': ['risk', 'hazard', 'incident'],
    'cac': ['credential', 'id', 'access'],
    'evidence': ['chain', 'custody', 'storage'],
    'haircut': ['grooming', 'appearance', 'hair'],
    'grooming': ['haircut', 'beard', 'mustache', 'appearance'],
    'beard': ['grooming', 'facial', 'hair'],
    'tattoo': ['appearance', 'uniform', 'grooming'],
    'uniform': ['appearance', 'dress', 'standards'],
    'pt': ['fitness', 'standards', 'physical'],
    'liberty': ['leave', 'off-duty', 'hours'],
    'overtime': ['timekeeping', 'civilian', 'payroll'],
    'desk': ['desk sergeant', 'front desk', 'watch'],
    'authorization': ['authority', 'approval', 'consent', 'warrant'],
    'authorize': ['authority', 'approval', 'consent'],
    'search': ['seizure', 'inspection', 'consent', 'gate'],
}

_ORDER_WEAK_TOKENS = {
    'order',
    'orders',
    'policy',
    'policies',
    'directive',
    'directives',
    'instruction',
    'instructions',
    'guidance',
    'guideline',
    'guidelines',
    'memorandum',
    'memorandums',
    'memo',
    'memos',
    'standard',
    'standards',
    'procedure',
    'procedures',
}


def _expand_search_terms(search_term):
    tokens = _tokenize_search_text(search_term)
    expanded = list(tokens)
    for token in tokens:
        if token in _ORDER_WEAK_TOKENS:
            continue
        for mapped in _ORDER_SYNONYM_MAP.get(token, []):
            if mapped not in expanded:
                expanded.append(mapped)
    phrase = (search_term or '').strip().lower()
    if 'traffic stop' in phrase:
        expanded.extend([item for item in ['vehicle', 'enforcement', 'citation'] if item not in expanded])
    if 'watch commander' in phrase:
        expanded.extend([item for item in ['supervisor', 'duty', 'command'] if item not in expanded])
    if 'lost cac' in phrase or 'lost id' in phrase:
        expanded.extend([item for item in ['credential', 'access', 'security'] if item not in expanded])
    return expanded


def _primary_search_terms(search_term):
    tokens = _tokenize_search_text(search_term)
    focused = [token for token in tokens if token not in _ORDER_WEAK_TOKENS]
    return focused or tokens


def _adjacent_query_phrases(tokens):
    phrases = []
    for idx in range(max(0, len(tokens) - 1)):
        phrase = f'{tokens[idx]} {tokens[idx + 1]}'.strip()
        if phrase and phrase not in phrases:
            phrases.append(phrase)
    return phrases


def _matched_primary_concepts(lower_sections, search_term):
    primary_tokens = _primary_search_terms(search_term)
    if not primary_tokens:
        return set()
    matched = set()
    for token in primary_tokens:
        concept_terms = [token]
        for mapped in _ORDER_SYNONYM_MAP.get(token, []):
            if mapped not in concept_terms:
                concept_terms.append(mapped)
        if any(term in text for text in lower_sections.values() for term in concept_terms):
            matched.add(token)
    return matched


def _snippet_for_query(text, search_term):
    if not text:
        return ''
    tokens = _tokenize_search_text(search_term)[:4]
    lower_text = text.lower()
    start = 0
    for token in tokens:
        pos = lower_text.find(token)
        if pos >= 0:
            start = max(0, pos - 80)
            break
    snippet = text[start:start + 220].replace('\n', ' ').strip()
    return snippet


def _searchable_sections(document):
    sections = []
    for label, value in (
        ('title', document.title or ''),
        ('order_number', document.order_number or ''),
        ('memo_number', document.memo_number or ''),
        ('summary', document.summary or ''),
        ('topic_tags', document.topic_tags or ''),
        ('audience_tags', document.audience_tags or ''),
        ('source_group', document.source_group or ''),
        ('issuing_authority', document.issuing_authority or ''),
    ):
        cleaned = (value or '').strip()
        if cleaned:
            sections.append((label, cleaned))

    raw_text = (document.extracted_text or '').replace('\r', '\n')
    lines = [line.strip() for line in raw_text.splitlines() if line.strip()]
    body_lines = [
        line for line in lines
        if not re.match(r'^(title|category|order number|memo number|issuing authority|summary|aliases|audience|topics):', line, flags=re.I)
    ]
    body_text = '\n'.join(body_lines).strip()
    if body_text:
        sections.append(('body', body_text))
    return sections


def _document_body_text(document):
    return next((value for label, value in _searchable_sections(document) if label == 'body'), '')


def _best_excerpt_with_reference(text, search_term):
    if not text:
        return '', ''
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if not lines:
        return _snippet_for_query(text, search_term), ''
    tokens = _expand_search_terms(search_term)[:10]
    match_index = -1
    for idx, line in enumerate(lines):
        lower_line = line.lower()
        if any(token in lower_line for token in tokens):
            match_index = idx
            break
    if match_index < 0:
        return _snippet_for_query(text, search_term), ''
    start = max(0, match_index - 1)
    end = min(len(lines), match_index + 2)
    excerpt = ' '.join(lines[start:end])[:420]
    reference = f"line {match_index + 1}"
    for back in range(match_index, max(-1, match_index - 6), -1):
        candidate = lines[back]
        if candidate.lower().startswith(('section ', 'para ', 'paragraph ', 'chapter ', 'article ', 'part ')):
            reference = candidate[:120]
            break
    return excerpt, reference


def _reader_blocks(text):
    blocks = []
    for raw_line in (text or '').replace('\r', '\n').splitlines():
        line = raw_line.strip()
        if not line:
            continue
        is_heading = bool(
            re.match(r'^(section|chapter|para|paragraph|article|part)\b', line, flags=re.I)
            or (len(line) <= 90 and line == line.title() and len(line.split()) <= 8)
        )
        blocks.append({'text': line, 'kind': 'heading' if is_heading else 'paragraph'})
    return blocks


def _reader_context(text, search_term, before=1, after=1, extra_before=3, extra_after=3):
    blocks = _reader_blocks(text)
    if not blocks:
        return {
            'focus_blocks': [],
            'leading_blocks': [],
            'trailing_blocks': [],
            'extra_blocks': [],
            'reference': '',
            'focus_index': 0,
            'all_blocks': [],
        }

    direct_tokens = _tokenize_search_text(search_term)[:6]
    tokens = direct_tokens or _expand_search_terms(search_term)[:12]
    focus_index = 0
    for token_set in (direct_tokens, _expand_search_terms(search_term)[:12]):
        if not token_set:
            continue
        for idx, block in enumerate(blocks):
            lower = block['text'].lower()
            if any(token in lower for token in token_set):
                focus_index = idx
                break
        else:
            continue
        break

    reference = ''
    for idx in range(focus_index, -1, -1):
        if blocks[idx]['kind'] == 'heading':
            reference = blocks[idx]['text']
            break
    if not reference:
        reference = f"paragraph {focus_index + 1}"

    lead_start = max(0, focus_index - before)
    trail_end = min(len(blocks), focus_index + after + 1)
    extra_start = max(0, focus_index - extra_before)
    extra_end = min(len(blocks), focus_index + extra_after + 1)

    leading_blocks = blocks[lead_start:focus_index]
    focus_blocks = [blocks[focus_index]]
    trailing_blocks = blocks[focus_index + 1:trail_end]
    extra_blocks = blocks[extra_start:lead_start] + blocks[trail_end:extra_end]

    return {
        'focus_blocks': focus_blocks,
        'leading_blocks': leading_blocks,
        'trailing_blocks': trailing_blocks,
        'extra_blocks': extra_blocks,
        'reference': reference,
        'focus_index': focus_index,
        'all_blocks': blocks,
    }


def _document_currentity_score(document):
    score = 0
    if document.is_active:
        score += 12
    stamp = document.revision_date or document.issue_date or document.last_indexed_at or document.uploaded_at
    if stamp:
        age_days = max(0, (datetime.now(timezone.utc).replace(tzinfo=None) - stamp).days)
        if age_days <= 365:
            score += 10
        elif age_days <= 365 * 2:
            score += 6
        elif age_days <= 365 * 4:
            score += 2
    if document.superseded_by_id:
        score -= 14
    return score


def _order_search_rank(document, search_term, extra_terms=None, topic_hints=None, audience_hints=None):
    if not search_term:
        return 10 + _document_currentity_score(document), ['No search phrase applied'], '', ''
    tokens = _tokenize_search_text(search_term)
    primary_tokens = _primary_search_terms(search_term)
    expanded_tokens = _expand_search_terms(search_term)
    expanded_tokens.extend(
        token
        for token in _tokenize_search_text(' '.join(extra_terms or ()))
        if token not in expanded_tokens
    )
    searchable_sections = _searchable_sections(document)
    haystack_fields = [value for _, value in searchable_sections] + [
        document.category or '',
        document.version_label or '',
        document.source_version or '',
    ]
    haystack = ' '.join(haystack_fields).lower()
    if not expanded_tokens:
        return 5 + _document_currentity_score(document), ['Search phrase too short'], '', ''

    score = 0
    reasons = []
    section_weights = {
        'title': 34,
        'order_number': 26,
        'memo_number': 26,
        'summary': 22,
        'topic_tags': 20,
        'audience_tags': 12,
        'source_group': 10,
        'issuing_authority': 8,
        'body': 14,
    }
    token_hits = set()
    primary_hits = set()
    section_primary_hits = {}
    phrase = search_term.lower()
    primary_phrase = ' '.join(primary_tokens).strip()
    adjacent_phrases = _adjacent_query_phrases(primary_tokens)

    for token in expanded_tokens:
        if token in primary_tokens:
            token_weight_bonus = 4
        elif token in tokens:
            token_weight_bonus = -8 if token in _ORDER_WEAK_TOKENS else 0
        else:
            token_weight_bonus = -12
        for label, section_text in searchable_sections:
            lower_section = section_text.lower()
            if token in lower_section:
                score += max(2, section_weights.get(label, 6) + token_weight_bonus)
                reasons.append(f"{label}:{token}")
                token_hits.add(token)
                if token in primary_tokens:
                    primary_hits.add(token)
                    section_primary_hits.setdefault(label, set()).add(token)
                if phrase and phrase in lower_section:
                    score += 12 if label in {'title', 'summary', 'body'} else 6
                    reasons.append(f"phrase in {label}")
                break

    if tokens:
        coverage = len({token for token in tokens if token in token_hits})
        score += coverage * 8
        if coverage == len(tokens):
            score += 12
            reasons.append('all query terms matched')

    if primary_tokens:
        primary_coverage = len(primary_hits)
        score += primary_coverage * 11
        if primary_coverage == len(primary_tokens):
            score += 18
            reasons.append('all primary terms matched')
        elif primary_coverage == 0:
            score -= 18
        elif len(primary_tokens) >= 2 and primary_coverage == 1:
            score -= 4

    if phrase in (document.title or '').lower():
        score += 24
        reasons.append('exact phrase in title')
    if phrase in (document.summary or '').lower():
        score += 18
        reasons.append('exact phrase in summary')
    body_text = next((value for label, value in searchable_sections if label == 'body'), '')
    if phrase and phrase in body_text.lower():
        score += 16
        reasons.append('exact phrase in relevant text')

    section_lookup = {label: value.lower() for label, value in searchable_sections}
    if primary_phrase:
        for label in ('title', 'summary', 'body'):
            if primary_phrase in section_lookup.get(label, ''):
                score += 18 if label == 'title' else 10
                reasons.append(f'focused phrase in {label}')
                break
    for label in ('title', 'summary', 'body'):
        hits = section_primary_hits.get(label, set())
        if len(hits) >= 2:
            score += 18 if label == 'title' else 10
            reasons.append(f'primary terms aligned in {label}')
        if any(query_phrase in section_lookup.get(label, '') for query_phrase in adjacent_phrases):
            score += 12 if label == 'title' else 8
            reasons.append(f'query phrase aligned in {label}')

    if any(token in (document.topic_tags or '').lower() for token in expanded_tokens):
        score += 10
        reasons.append('topic tags aligned')

    if any(token in (document.audience_tags or '').lower() for token in expanded_tokens):
        score += 4
        reasons.append('audience tags aligned')

    for hint in _dedupe_search_phrases(topic_hints or [], limit=4):
        if hint.lower() in ((document.topic_tags or '') + ' ' + (document.summary or '') + ' ' + body_text).lower():
            score += 10
            reasons.append(f'ai topic:{hint}')

    for hint in _dedupe_search_phrases(audience_hints or [], limit=3):
        if hint.lower() in ((document.audience_tags or '') + ' ' + (document.summary or '')).lower():
            score += 8
            reasons.append(f'ai audience:{hint}')

    score += _document_currentity_score(document)
    snippet, reference = _best_excerpt_with_reference(body_text or document.summary or '', search_term)
    if not snippet:
        snippet = _snippet_for_query(body_text or document.summary or '', search_term)
    return score, reasons[:6], snippet, reference


def _is_relevant_order_match(document, search_term, score):
    if not search_term:
        return True
    tokens = _tokenize_search_text(search_term)
    primary_tokens = _primary_search_terms(search_term)
    if not tokens:
        return False
    searchable_sections = _searchable_sections(document)
    lower_sections = {label: value.lower() for label, value in searchable_sections}
    phrase = search_term.lower().strip()
    if phrase and any(phrase in text for text in lower_sections.values()):
        return True

    matched_primary_concepts = _matched_primary_concepts(lower_sections, search_term)
    direct_hits = {
        token for token in tokens
        if any(token in text for text in lower_sections.values())
    }
    primary_hits = {
        token for token in primary_tokens
        if any(token in text for text in lower_sections.values())
    }
    strong_direct_hit = any(
        token in lower_sections.get('title', '')
        or token in lower_sections.get('summary', '')
        or token in lower_sections.get('body', '')
        or token in lower_sections.get('topic_tags', '')
        or token in lower_sections.get('order_number', '')
        or token in lower_sections.get('memo_number', '')
        for token in (matched_primary_concepts or primary_hits or direct_hits)
    )

    if len(primary_tokens) == 1:
        return bool(matched_primary_concepts) and score >= 18
    if len(primary_tokens) == 2:
        return len(matched_primary_concepts) == 2 and score >= 24
    if len(matched_primary_concepts) >= 2:
        return score >= 24
    if len(matched_primary_concepts) == 1 and strong_direct_hit and score >= 54:
        return True
    return False


def _humanize_order_match_reason(reason):
    raw = str(reason or '').strip()
    if not raw:
        return ''
    replacements = {
        'phrase in title': 'matched the full phrase in the title',
        'phrase in summary': 'matched the full phrase in the summary',
        'phrase in body': 'matched the full phrase in the source text',
        'exact phrase in title': 'matched the exact phrase in the title',
        'exact phrase in summary': 'matched the exact phrase in the summary',
        'exact phrase in relevant text': 'matched the exact phrase in the source text',
        'all query terms matched': 'covered all original search terms',
        'all primary terms matched': 'covered all focused search terms',
        'focused phrase in title': 'kept the focused phrase together in the title',
        'focused phrase in summary': 'kept the focused phrase together in the summary',
        'focused phrase in body': 'kept the focused phrase together in the source text',
        'primary terms aligned in title': 'kept the key terms together in the title',
        'primary terms aligned in summary': 'kept the key terms together in the summary',
        'primary terms aligned in body': 'kept the key terms together in the source text',
        'query phrase aligned in title': 'matched a key phrase in the title',
        'query phrase aligned in summary': 'matched a key phrase in the summary',
        'query phrase aligned in body': 'matched a key phrase in the source text',
        'topic tags aligned': 'aligned with the document topic tags',
        'audience tags aligned': 'aligned with the intended audience',
    }
    if raw in replacements:
        return replacements[raw]
    if ':' not in raw:
        return raw.replace('_', ' ').capitalize()

    label, value = raw.split(':', 1)
    friendly_label = {
        'title': 'title',
        'summary': 'summary',
        'body': 'source text',
        'topic_tags': 'topic tags',
        'audience_tags': 'audience tags',
        'source_group': 'source group',
        'issuing_authority': 'issuing authority',
        'order_number': 'order number',
        'memo_number': 'memo number',
        'ai topic': 'topic hint',
        'ai audience': 'audience hint',
    }.get(label, label.replace('_', ' '))
    return f'matched "{value}" in the {friendly_label}'


def _display_order_match_reasons(reasons, limit=4):
    display = []
    seen = set()
    for reason in reasons or []:
        humanized = _humanize_order_match_reason(reason)
        if not humanized:
            continue
        lowered = humanized.lower()
        if lowered in seen:
            continue
        seen.add(lowered)
        display.append(humanized)
        if len(display) >= limit:
            break
    return display


def _filtered_orders(search_term='', category_filter='', status_filter='ACTIVE', source_type_filter='', year_filter='', topic_filter='', source_filter=''):
    query = OrderDocument.query
    if category_filter:
        query = query.filter(OrderDocument.category.ilike(category_filter))
    if source_type_filter:
        query = query.filter(OrderDocument.source_type == source_type_filter)
    if year_filter and year_filter.isdigit():
        year_start = datetime(int(year_filter), 1, 1)
        year_end = datetime(int(year_filter) + 1, 1, 1)
        query = query.filter(
            db.or_(
                db.and_(OrderDocument.issue_date.isnot(None), OrderDocument.issue_date >= year_start, OrderDocument.issue_date < year_end),
                db.and_(OrderDocument.revision_date.isnot(None), OrderDocument.revision_date >= year_start, OrderDocument.revision_date < year_end),
                db.and_(OrderDocument.uploaded_at.isnot(None), OrderDocument.uploaded_at >= year_start, OrderDocument.uploaded_at < year_end),
            )
        )
    if topic_filter:
        topic_like = f"%{topic_filter}%"
        query = query.filter(
            db.or_(
                OrderDocument.topic_tags.ilike(topic_like),
                OrderDocument.summary.ilike(topic_like),
                OrderDocument.extracted_text.ilike(topic_like),
            )
        )
    if source_filter:
        source_like = f"%{source_filter}%"
        query = query.filter(OrderDocument.source_group.ilike(source_like))
    status_filter = (status_filter or 'ACTIVE').strip().upper()
    if status_filter == 'ACTIVE':
        query = query.filter(OrderDocument.is_active.is_(True))
    elif status_filter == 'INACTIVE':
        query = query.filter(OrderDocument.is_active.is_(False))
    documents = query.order_by(
        OrderDocument.revision_date.desc().nullslast(),
        OrderDocument.uploaded_at.desc(),
        OrderDocument.id.desc(),
    ).all()
    approved_config = _load_approved_sources_config()
    ranked = []
    for document in documents:
        source_meta = _document_source_metadata(document, approved_config)
        if not source_meta['approved']:
            continue
        score, reasons, snippet, reference = _order_search_rank(document, search_term)
        if search_term and (score <= 0 or not _is_relevant_order_match(document, search_term, score)):
            continue
        document.search_score = score
        document.search_confidence = max(15, min(99, 20 + score))
        document.match_reasons_raw = reasons
        document.match_reasons = _display_order_match_reasons(reasons)
        document.match_snippet = snippet
        document.match_reference = reference
        document.audience_label = (document.audience_tags or '').split(',')[0].strip() if (document.audience_tags or '').strip() else ''
        document.download_available = _document_has_download(document)
        document.source_origin_label = source_meta['origin_label']
        document.source_origin_type = source_meta['origin_type']
        ranked.append(document)
    ranked.sort(key=lambda item: (getattr(item, 'search_score', 0), item.uploaded_at or datetime.min), reverse=True)
    return ranked


def _civilian_search_recommendations(search_term, documents):
    normalized = (search_term or '').strip().lower()
    suggestions = []
    if not normalized:
        return suggestions
    if any(token in normalized for token in ('leave', 'vacation', 'sick', 'absence', 'time off')):
        suggestions.append('Try personnel, attendance, or leave-related orders first.')
    if any(token in normalized for token in ('badge', 'access', 'credential', 'gate', 'visitor')):
        suggestions.append('Check security, access-control, or base access orders.')
    if any(token in normalized for token in ('vehicle', 'parking', 'traffic', 'accident')):
        suggestions.append('Look for traffic, vehicle operations, or parking directives.')
    if any(token in normalized for token in ('report', 'form', 'submit', 'paperwork')):
        suggestions.append('Search document titles for reporting, forms, or administrative instructions.')
    if documents and not suggestions:
        suggestions.append('Matching orders were found below based on title, summary, category, or version text.')
    elif not documents:
        suggestions.append('No direct match found. Try broader words like access, training, leave, security, or parking.')
    return suggestions


def _safe_ai_guidance(prompt):
    if not (prompt or '').strip():
        return ''
    answer = ask_openai(prompt, os.environ.get('OPENAI_API_KEY'))
    if not answer:
        return ''
    if is_ai_unavailable_message(answer):
        return ''
    return answer


def _ai_order_search_strategy(search_term, category_filter='', source_type_filter='', current_documents=None):
    clean_search = (search_term or '').strip()
    if not clean_search:
        return _empty_order_ai_strategy()
    current_documents = current_documents or []
    filter_note = ', '.join(
        part for part in (
            f'category={category_filter}' if category_filter else '',
            f'source_type={source_type_filter}' if source_type_filter else '',
        ) if part
    ) or 'no extra filters'
    candidate_summary = _order_result_summary(current_documents, limit=6) if current_documents else 'No strong approved matches yet.'
    prompt = (
        "You are improving retrieval for an MCPD internal order and memorandum search. "
        "You are NOT allowed to invent document titles, order numbers, or policy conclusions. "
        "Your job is only to generate grounded local-search hints for approved documents already in the library. "
        "Return STRICT JSON with keys: priority_terms (array), query_variants (array), topic_hints (array), audience_hints (array), officer_brief (string). "
        "Keep all phrases short and practical. Use no markdown.\n\n"
        f"User request: {clean_search}\n"
        f"Active filters: {filter_note}\n"
        f"Current approved candidates:\n{candidate_summary}"
    )
    answer = _safe_ai_guidance(prompt)
    if not answer:
        return _empty_order_ai_strategy()
    return _normalize_ai_strategy(_extract_json_object(answer))


def _merge_ranked_documents(*document_groups):
    merged = {}
    for group in document_groups:
        for document in group or []:
            existing = merged.get(document.id)
            existing_score = getattr(existing, 'search_score', -1) if existing is not None else -1
            current_score = getattr(document, 'search_score', 0)
            if existing is None or current_score > existing_score:
                merged[document.id] = document
    return sorted(
        merged.values(),
        key=lambda item: (getattr(item, 'search_score', 0), item.uploaded_at or datetime.min),
        reverse=True,
    )


def _apply_ai_order_boost(documents, strategy):
    if not documents:
        return []
    strategy = strategy or _empty_order_ai_strategy()
    priority_terms = _dedupe_search_phrases(strategy.get('priority_terms') or [], limit=8)
    topic_hints = _dedupe_search_phrases(strategy.get('topic_hints') or [], limit=6)
    audience_hints = _dedupe_search_phrases(strategy.get('audience_hints') or [], limit=4)
    for document in documents:
        haystack = ' '.join(
            filter(
                None,
                [
                    document.title or '',
                    document.summary or '',
                    document.topic_tags or '',
                    document.audience_tags or '',
                    document.extracted_text or '',
                ],
            )
        ).lower()
        bonus = 0
        if any(term.lower() in haystack for term in priority_terms):
            bonus += 12
        if any(term.lower() in haystack for term in topic_hints):
            bonus += 10
        if any(term.lower() in haystack for term in audience_hints):
            bonus += 8
        if bonus:
            document.search_score = getattr(document, 'search_score', 0) + bonus
            document.search_confidence = max(15, min(99, 20 + document.search_score))
            existing_reasons = list(getattr(document, 'match_reasons', []) or [])
            if 'AI retrieval alignment' not in existing_reasons:
                existing_reasons.append('AI retrieval alignment')
            document.match_reasons = existing_reasons[:8]
    return sorted(
        documents,
        key=lambda item: (getattr(item, 'search_score', 0), item.uploaded_at or datetime.min),
        reverse=True,
    )


def search_orders_with_ai_assist(search_term='', category_filter='', status_filter='ACTIVE', source_type_filter='', year_filter='', topic_filter='', source_filter=''):
    documents = _filtered_orders(
        search_term,
        category_filter,
        status_filter,
        source_type_filter=source_type_filter,
        year_filter=year_filter,
        topic_filter=topic_filter,
        source_filter=source_filter,
    )
    strategy = _empty_order_ai_strategy()
    clean_search = (search_term or '').strip()
    if not clean_search or not current_app.config.get('ORDERS_AI_ASSIST_ENABLED', False):
        return documents, strategy

    strategy = _ai_order_search_strategy(
        clean_search,
        category_filter=category_filter,
        source_type_filter=source_type_filter,
        current_documents=documents,
    )
    if not any(strategy.get(key) for key in ('priority_terms', 'query_variants', 'topic_hints', 'audience_hints', 'officer_brief')):
        return documents, strategy

    variant_queries = _dedupe_search_phrases(
        list(strategy.get('query_variants') or [])
        + ([f"{clean_search} {' '.join((strategy.get('priority_terms') or [])[:4])}".strip()] if strategy.get('priority_terms') else []),
        limit=4,
    )
    variant_documents = []
    for variant in variant_queries:
        if variant.lower() == clean_search.lower():
            continue
        variant_documents.append(
            _filtered_orders(
                variant,
                category_filter,
                status_filter,
                source_type_filter=source_type_filter,
                year_filter=year_filter,
                topic_filter=topic_filter,
                source_filter=source_filter,
            )
        )
    merged = _merge_ranked_documents(documents, *variant_documents)
    boosted = _apply_ai_order_boost(merged, strategy)
    return boosted, strategy


def _order_result_summary(documents, limit=6):
    lines = []
    for document in documents[:limit]:
        lines.append(
            f"- {document.title} | category={document.category or 'General'} | summary={document.summary or 'No summary'}"
        )
    return '\n'.join(lines)


def _local_order_brief(documents):
    if not documents:
        return ''
    top = documents[0]
    reasons = ', '.join((getattr(top, 'match_reasons', []) or [])[:3])
    if reasons:
        return f"Open {top.title} first. It ranked highest based on {reasons}."
    return f"Open {top.title} first. It is the strongest approved policy match in the current library."


def _approved_filter_facets():
    approved_documents = _approved_documents()
    categories = sorted({(doc.category or '').strip() for doc in approved_documents if (doc.category or '').strip()}, key=str.lower)
    source_groups = sorted({(doc.source_group or '').strip() for doc in approved_documents if (doc.source_group or '').strip()}, key=str.lower)
    years = sorted(
        {
            (doc.issue_date or doc.revision_date or doc.uploaded_at).year
            for doc in approved_documents
            if (doc.issue_date or doc.revision_date or doc.uploaded_at)
        },
        reverse=True,
    )
    topic_values = sorted(
        {
            item.strip()
            for doc in approved_documents
            for item in (doc.topic_tags or '').split(',')
            if item.strip()
        }
    )
    return categories, source_groups, years, topic_values


@bp.route('/orders', methods=['GET', 'POST'])
@login_required
def library():
    if not current_user.can_manage_team():
        return redirect(url_for('orders.reference_search', **request.values.to_dict(flat=True)))
    if request.method == 'POST':
        seed_action = (request.form.get('action') or '').strip().lower()
        if seed_action == 'seed_library':
            if not current_user.can_manage_team():
                abort(403)
            inserted, skipped = _seed_orders_library()
            if inserted:
                flash(f'Starter order library loaded: {inserted} added, {skipped} skipped.', 'success')
            else:
                flash('Starter library already present or seed file missing.', 'error')
            return redirect(url_for('orders.library'))
        if seed_action == 'ingest_official_sources':
            if not current_user.can_manage_team():
                abort(403)
            max_new = request.form.get('max_new', type=int) or 20
            result = run_official_orders_ingestion(max_new=max(1, min(max_new, 100)), fetch_limit=140)
            if result.get('ok'):
                flash(
                    f"Official ingestion complete: scanned {result['scanned']}, inserted {result['inserted']}, updated {result['updated']}.",
                    'success',
                )
            else:
                flash('Official ingestion completed with errors. Check server logs/state.', 'warning')
            return redirect(url_for('orders.library'))
        if seed_action == 'reindex_library':
            if not current_user.can_manage_team():
                abort(403)
            docs = OrderDocument.query.order_by(OrderDocument.id.desc()).all()
            updated = 0
            for doc in docs:
                if doc.file_path and os.path.exists(doc.file_path):
                    extracted = _extract_text_from_file(doc.file_path)
                    if extracted and extracted != (doc.extracted_text or ''):
                        doc.extracted_text = extracted
                        doc.parser_confidence = 0.9
                        doc.last_indexed_at = datetime.utcnow()
                        updated += 1
            if updated:
                db.session.commit()
                flash(f'Reindex complete: {updated} document(s) refreshed.', 'success')
            else:
                flash('Reindex complete: no updates required.', 'success')
            return redirect(url_for('orders.library'))
        if seed_action == 'reindex_approved_sources':
            if not current_user.can_manage_team():
                abort(403)
            docs = _approved_documents()
            updated = 0
            for doc in docs:
                if doc.file_path and os.path.exists(doc.file_path):
                    extracted = _extract_text_from_file(doc.file_path)
                    if extracted and extracted != (doc.extracted_text or ''):
                        doc.extracted_text = extracted
                        doc.parser_confidence = 0.9
                        doc.last_indexed_at = datetime.utcnow()
                        updated += 1
            if updated:
                db.session.commit()
                flash(f'Approved-source reindex complete: {updated} document(s) refreshed.', 'success')
            else:
                flash('Approved-source reindex complete: no updates required.', 'success')
            return redirect(url_for('orders.orders_source_admin'))

        if not current_user.can_manage_team():
            abort(403)

        upload_file = request.files.get('order_file')
        title = (request.form.get('title') or '').strip()
        category = (request.form.get('category') or '').strip() or None
        source_type = (request.form.get('source_type') or '').strip().upper() or 'LOCAL_DOCUMENT'
        source_group = (request.form.get('source_group') or '').strip() or None
        order_number = (request.form.get('order_number') or '').strip() or None
        memo_number = (request.form.get('memo_number') or '').strip() or None
        issuing_authority = (request.form.get('issuing_authority') or '').strip() or None
        issue_date = _parse_date(request.form.get('issue_date'))
        revision_date = _parse_date(request.form.get('revision_date'))
        source_version = (request.form.get('source_version') or '').strip() or None
        audience_tags = (request.form.get('audience_tags') or '').strip() or None
        topic_tags = (request.form.get('topic_tags') or '').strip() or None
        superseded_by_id = request.form.get('superseded_by_id', type=int)
        version_label = (request.form.get('version_label') or '').strip() or None
        summary = (request.form.get('summary') or '').strip() or None

        if not upload_file or not upload_file.filename:
            flash('Select a file to upload.', 'error')
            return redirect(url_for('orders.library'))
        if not title:
            flash('Title is required.', 'error')
            return redirect(url_for('orders.library'))

        original_name = secure_filename(upload_file.filename)
        extension = os.path.splitext(original_name)[1].lower()
        if extension not in ALLOWED_EXTENSIONS:
            flash('Upload a PDF, Word, text, or RTF document.', 'error')
            return redirect(url_for('orders.library'))

        stamped_name = f"{int(datetime.utcnow().timestamp())}-{original_name}"
        save_path = os.path.join(_orders_directory(), stamped_name)
        upload_file.save(save_path)
        extracted_text = _extract_text_from_file(save_path)
        parser_confidence = 0.9 if extracted_text else 0.35
        if source_type not in SOURCE_TYPE_OPTIONS:
            source_type = 'LOCAL_DOCUMENT'

        document = OrderDocument(
            title=title,
            category=category,
            source_type=source_type,
            source_group=source_group,
            order_number=order_number,
            memo_number=memo_number,
            issuing_authority=issuing_authority,
            issue_date=issue_date,
            revision_date=revision_date,
            source_version=source_version,
            audience_tags=audience_tags,
            topic_tags=topic_tags,
            version_label=version_label,
            summary=summary,
            extracted_text=extracted_text,
            parser_confidence=parser_confidence,
            file_path=save_path,
            uploaded_by=current_user.id,
            superseded_by_id=superseded_by_id,
            last_indexed_at=datetime.utcnow(),
        )
        db.session.add(document)
        db.session.commit()
        flash('Order uploaded.', 'success')
        return redirect(url_for('orders.library'))

    search_term = (request.args.get('q') or '').strip()
    category_filter = (request.args.get('category') or '').strip()
    status_filter = (request.args.get('status') or 'ACTIVE').strip().upper()
    source_type_filter = (request.args.get('source_type') or '').strip().upper()
    year_filter = (request.args.get('year') or '').strip()
    topic_filter = (request.args.get('topic') or '').strip()
    source_filter = (request.args.get('source') or '').strip()

    documents, ai_strategy = search_orders_with_ai_assist(
        search_term,
        category_filter,
        status_filter,
        source_type_filter=source_type_filter,
        year_filter=year_filter,
        topic_filter=topic_filter,
        source_filter=source_filter,
    )
    ai_guidance = (ai_strategy.get('officer_brief') or '').strip() or _local_order_brief(documents)
    categories, source_groups, years, topic_values = _approved_filter_facets()

    return render_template(
        'orders.html',
        user=current_user,
        documents=documents,
        categories=categories,
        order_search_term=search_term,
        order_category_filter=category_filter,
        order_status_filter=status_filter,
        order_source_type_filter=source_type_filter,
        order_year_filter=year_filter,
        order_topic_filter=topic_filter,
        order_source_filter=source_filter,
        can_upload_orders=current_user.can_manage_team(),
        source_type_options=SOURCE_TYPE_OPTIONS,
        source_groups=source_groups,
        years=years,
        topic_values=topic_values,
        known_documents=OrderDocument.query.order_by(OrderDocument.uploaded_at.desc()).limit(150).all(),
        display_dt=_safe_display_dt,
        ai_guidance=ai_guidance,
        ai_priority_terms=ai_strategy.get('priority_terms', ()),
        ai_query_variants=ai_strategy.get('query_variants', ()),
        ai_topic_hints=ai_strategy.get('topic_hints', ()),
        ai_audience_hints=ai_strategy.get('audience_hints', ()),
        orders_seed_available=bool(_load_seed_orders()),
        ingestion_state=get_ingestion_state(),
    )


@bp.route('/orders/reference', methods=['GET', 'POST'])
@login_required
def reference_search():
    _ensure_orders_library_seeded()
    if request.method == 'POST':
        return redirect(url_for('orders.library'))

    search_term = (request.args.get('q') or '').strip()
    status_filter = (request.args.get('status') or 'ACTIVE').strip().upper()
    category_filter = (request.args.get('category') or '').strip()
    source_type_filter = (request.args.get('source_type') or '').strip().upper()
    year_filter = (request.args.get('year') or '').strip()
    topic_filter = (request.args.get('topic') or '').strip()
    source_filter = (request.args.get('source') or '').strip()
    documents, ai_strategy = search_orders_with_ai_assist(
        search_term,
        category_filter,
        status_filter,
        source_type_filter=source_type_filter,
        year_filter=year_filter,
        topic_filter=topic_filter,
        source_filter=source_filter,
    )
    categories, source_groups, years, topic_values = _approved_filter_facets()
    ai_guidance = (ai_strategy.get('officer_brief') or '').strip() or _local_order_brief(documents)
    lead_document = documents[0] if documents else None
    secondary_documents = documents[1:] if len(documents) > 1 else []
    active_filters = []
    if status_filter and status_filter != 'ALL':
        active_filters.append({'label': 'Status', 'value': 'Active only' if status_filter == 'ACTIVE' else 'Archived only'})
    if category_filter:
        active_filters.append({'label': 'Category', 'value': category_filter})
    if source_type_filter:
        active_filters.append({'label': 'Type', 'value': source_type_filter})
    if year_filter:
        active_filters.append({'label': 'Year', 'value': year_filter})
    if topic_filter:
        active_filters.append({'label': 'Topic', 'value': topic_filter})
    if source_filter:
        active_filters.append({'label': 'Source', 'value': source_filter})
    return render_template(
        'orders_reference.html',
        user=current_user,
        documents=documents,
        lead_document=lead_document,
        secondary_documents=secondary_documents,
        categories=categories,
        search_term=search_term,
        category_filter=category_filter,
        status_filter=status_filter,
        source_type_filter=source_type_filter,
        year_filter=year_filter,
        topic_filter=topic_filter,
        source_filter=source_filter,
        can_manage_orders=current_user.can_manage_team(),
        source_type_options=SOURCE_TYPE_OPTIONS,
        source_groups=source_groups,
        years=years,
        topic_values=topic_values,
        display_dt=_safe_display_dt,
        ai_guidance=ai_guidance,
        ai_priority_terms=ai_strategy.get('priority_terms', ()),
        ai_query_variants=ai_strategy.get('query_variants', ()),
        ai_topic_hints=ai_strategy.get('topic_hints', ()),
        ai_audience_hints=ai_strategy.get('audience_hints', ()),
        active_filters=active_filters,
    )


@bp.route('/orders/admin/sources', methods=['GET', 'POST'])
@login_required
def orders_source_admin():
    if not current_user.can_manage_team():
        abort(403)

    config = _load_approved_sources_config()
    if request.method == 'POST':
        action = (request.form.get('action') or '').strip()
        if action == 'add_domain':
            domain = _normalized_host(request.form.get('domain') or '')
            if not domain:
                flash('Enter a valid domain to allow.', 'warning')
            elif domain not in {_normalized_host(item) for item in (config.get('approved_domains') or [])}:
                config['approved_domains'] = list(config.get('approved_domains') or []) + [domain]
                _save_approved_sources_config(config)
                flash(f'Approved domain added: {domain}', 'success')
            else:
                flash('Domain is already allowlisted.', 'info')
            return redirect(url_for('orders.orders_source_admin'))
        if action == 'remove_domain':
            domain = _normalized_host(request.form.get('domain') or '')
            config['approved_domains'] = [item for item in (config.get('approved_domains') or []) if _normalized_host(item) != domain]
            _save_approved_sources_config(config)
            flash(f'Approved domain removed: {domain}', 'success')
            return redirect(url_for('orders.orders_source_admin'))
        if action == 'add_source_group':
            source_group = (request.form.get('source_group') or '').strip()
            if not source_group:
                flash('Enter a source group to allow.', 'warning')
            elif source_group not in (config.get('approved_source_groups') or []):
                config['approved_source_groups'] = list(config.get('approved_source_groups') or []) + [source_group]
                _save_approved_sources_config(config)
                flash(f'Approved source group added: {source_group}', 'success')
            else:
                flash('Source group is already allowlisted.', 'info')
            return redirect(url_for('orders.orders_source_admin'))
        if action == 'remove_source_group':
            source_group = (request.form.get('source_group') or '').strip()
            config['approved_source_groups'] = [item for item in (config.get('approved_source_groups') or []) if item != source_group]
            _save_approved_sources_config(config)
            flash(f'Approved source group removed: {source_group}', 'success')
            return redirect(url_for('orders.orders_source_admin'))
        if action == 'archive_unapproved':
            updated = 0
            for doc in OrderDocument.query.order_by(OrderDocument.id.desc()).all():
                if _document_source_metadata(doc, config)['approved']:
                    continue
                if doc.is_active:
                    doc.is_active = False
                    updated += 1
            if updated:
                db.session.commit()
            flash(f'Archived {updated} unapproved document(s).' if updated else 'No active unapproved documents needed archiving.', 'success')
            return redirect(url_for('orders.orders_source_admin'))
        if action == 'archive_doc':
            document = OrderDocument.query.get_or_404(request.form.get('document_id', type=int))
            document.is_active = False
            db.session.commit()
            flash(f'Archived "{document.title}".', 'success')
            return redirect(url_for('orders.orders_source_admin'))
        if action == 'reindex_approved':
            updated = 0
            for doc in _approved_documents():
                if doc.file_path and os.path.exists(doc.file_path):
                    extracted = _extract_text_from_file(doc.file_path)
                    if extracted and extracted != (doc.extracted_text or ''):
                        doc.extracted_text = extracted
                        doc.parser_confidence = 0.9
                        doc.last_indexed_at = datetime.utcnow()
                        updated += 1
            if updated:
                db.session.commit()
            flash(f'Approved-source reindex complete: {updated} document(s) refreshed.' if updated else 'Approved-source reindex complete: no updates required.', 'success')
            return redirect(url_for('orders.orders_source_admin'))

    documents = OrderDocument.query.order_by(OrderDocument.uploaded_at.desc(), OrderDocument.id.desc()).all()
    origin_rows = {}
    unapproved_documents = []
    approved_count = 0
    for document in documents:
        metadata = _document_source_metadata(document, config)
        key = metadata['origin_label']
        row = origin_rows.setdefault(
            key,
            {
                'label': metadata['origin_label'],
                'origin_type': metadata['origin_type'],
                'approved': metadata['approved'],
                'count': 0,
                'active_count': 0,
                'sample_title': document.title,
            },
        )
        row['count'] += 1
        if document.is_active:
            row['active_count'] += 1
        if metadata['approved']:
            approved_count += 1
        else:
            document.source_block_reason = metadata['blocked_reason']
            document.source_origin_label = metadata['origin_label']
            unapproved_documents.append(document)
    source_rows = sorted(origin_rows.values(), key=lambda item: (0 if item['approved'] else 1, item['label'].lower()))

    return render_template(
        'orders_admin_sources.html',
        user=current_user,
        approved_config=config,
        source_rows=source_rows,
        unapproved_documents=unapproved_documents[:150],
        approved_count=approved_count,
        total_documents=len(documents),
        display_dt=_safe_display_dt,
    )


@bp.route('/orders/ingest/run', methods=['POST'])
@login_required
def run_orders_ingestion():
    if not current_user.can_manage_team():
        abort(403)
    max_new = request.form.get('max_new', type=int) or request.args.get('max_new', type=int) or 20
    result = run_official_orders_ingestion(max_new=max(1, min(max_new, 100)), fetch_limit=140)
    if request.headers.get('X-Requested-With') == 'XMLHttpRequest' or request.accept_mimetypes.accept_json:
        return jsonify(result)
    if result.get('ok'):
        flash(
            f"Official ingestion complete: scanned {result['scanned']}, inserted {result['inserted']}, updated {result['updated']}.",
            'success',
        )
    else:
        flash('Official ingestion completed with errors. Check logs/state.', 'warning')
    return redirect(url_for('orders.library'))


@bp.route('/orders/ingest/status', methods=['GET'])
@login_required
def orders_ingest_status():
    if not current_user.can_manage_team():
        abort(403)
    return jsonify(get_ingestion_state())


@bp.route('/orders/civilian-help', methods=['GET'])
@login_required
def civilian_help():
    _ensure_orders_library_seeded()
    search_term = (request.args.get('q') or '').strip()
    documents = _filtered_orders(search_term, '', 'ACTIVE')
    suggestions = _civilian_search_recommendations(search_term, documents)
    ai_guidance = ''
    if search_term:
        ai_guidance = _safe_ai_guidance(
            "You are helping a civilian employee find the right MCPD document or directive. "
            "Based on the user request and the available matching documents, explain in plain language what they should open first and what kind of document they are probably looking for. "
            "Keep it simple, practical, and short.\n\n"
            f"User request: {search_term}\n\n"
            f"Matching documents:\n{_order_result_summary(documents)}"
        )
    return render_template(
        'civilian_help.html',
        user=current_user,
        documents=documents[:15],
        search_term=search_term,
        suggestions=suggestions,
        display_dt=_safe_display_dt,
        ai_guidance=ai_guidance,
    )


@bp.route('/orders/<int:order_id>/view', methods=['GET'])
@login_required
def view_order(order_id):
    _ensure_orders_library_seeded()
    document = OrderDocument.query.get_or_404(order_id)
    search_term = (request.args.get('q') or '').strip()
    download_path = _document_download_path(document)
    file_exists = bool(download_path)
    freshness_warning = ''
    if document.revision_date:
        age_days = (datetime.utcnow() - document.revision_date).days
        if age_days > 365 * 2:
            freshness_warning = 'source currency requires verification'
    elif document.issue_date:
        age_days = (datetime.utcnow() - document.issue_date).days
        if age_days > 365 * 2:
            freshness_warning = 'source currency requires verification'
    elif document.uploaded_at:
        age_days = (datetime.utcnow() - document.uploaded_at).days
        if age_days > 365 * 2:
            freshness_warning = 'source currency requires verification'

    newer_candidate = None
    if document.order_number or document.memo_number:
        newer_query = OrderDocument.query.filter(
            OrderDocument.id != document.id,
            OrderDocument.is_active.is_(True),
            db.or_(
                OrderDocument.order_number == document.order_number,
                OrderDocument.memo_number == document.memo_number,
            ),
        ).order_by(OrderDocument.revision_date.desc().nullslast(), OrderDocument.uploaded_at.desc())
        newer_candidate = newer_query.first()
    if newer_candidate:
        current_stamp = document.revision_date or document.uploaded_at or datetime.min
        newer_stamp = newer_candidate.revision_date or newer_candidate.uploaded_at or datetime.min
        if newer_stamp > current_stamp:
            freshness_warning = 'newer source may exist - verify current version'

    if not document.extracted_text and file_exists:
        document.extracted_text = _extract_text_from_file(download_path)
        document.last_indexed_at = datetime.utcnow()
        document.parser_confidence = 0.9 if document.extracted_text else 0.35
        db.session.commit()

    body_text = _document_body_text(document) or document.summary or ''
    reader_context = _reader_context(body_text, search_term or document.title, before=1, after=2, extra_before=3, extra_after=4)
    excerpt_reference = reader_context['reference']
    visible_blocks = reader_context['leading_blocks'] + reader_context['focus_blocks'] + reader_context['trailing_blocks']
    excerpt_text = '\n\n'.join(block['text'] for block in visible_blocks).strip()
    if not excerpt_text:
        excerpt_text = (document.summary or _snippet_for_query(body_text, search_term) or '').strip()

    simplify_warning = ''
    if document.parser_confidence and document.parser_confidence < 0.55:
        simplify_warning = 'relevant section may require manual review'
    partial_text_warning = ''
    if document.parser_confidence and document.parser_confidence < 0.75:
        partial_text_warning = 'Only partial extracted text may be available. Verify against the original source if needed.'

    return render_template(
        'orders_view.html',
        user=current_user,
        document=document,
        search_term=search_term,
        excerpt_text=excerpt_text,
        excerpt_reference=excerpt_reference,
        reader_context=reader_context,
        body_text=body_text,
        simplify_warning=simplify_warning,
        partial_text_warning=partial_text_warning,
        file_exists=file_exists,
        freshness_warning=freshness_warning,
        display_dt=_safe_display_dt,
    )


@bp.route('/orders/<int:order_id>/simplify', methods=['POST'])
@login_required
def simplify_order_excerpt(order_id):
    document = OrderDocument.query.get_or_404(order_id)
    payload = request.get_json(silent=True) or {}
    excerpt = (payload.get('excerpt') or '').strip()
    reference = (payload.get('reference') or '').strip()
    if not excerpt:
        excerpt = (document.summary or _snippet_for_query(document.extracted_text or '', '') or '').strip()
    if not excerpt:
        return jsonify({'ok': False, 'error': 'No source excerpt available to simplify.'}), 400

    prompt = (
        "You are generating a plain-English helper explanation for a policy/order reference. "
        "Do not invent rules. Keep the explanation concise and practical. "
        "Start with: 'Plain-English summary:'.\n\n"
        f"Document title: {document.title}\n"
        f"Citation/order: {document.order_number or document.memo_number or '-'}\n"
        f"Reference: {reference or 'not specified'}\n"
        f"Official excerpt:\n{excerpt}\n"
    )
    simplified = _safe_ai_guidance(prompt)
    if not simplified:
        simplified = 'Plain-English summary unavailable right now. Verify using the official source text.'
    return jsonify({'ok': True, 'summary': simplified})


@bp.route('/orders/<int:order_id>/download', methods=['GET'])
@login_required
def download_order(order_id):
    document = OrderDocument.query.get_or_404(order_id)
    download_path = _document_download_path(document)
    if not download_path:
        abort(404)
    return send_file(download_path, as_attachment=True, download_name=os.path.basename(download_path))


@bp.route('/orders/<int:order_id>/toggle', methods=['POST'])
@login_required
def toggle_order(order_id):
    if not current_user.can_manage_team():
        abort(403)

    document = OrderDocument.query.get_or_404(order_id)
    document.is_active = not document.is_active
    db.session.commit()
    flash(
        'Order activated.' if document.is_active else 'Order archived.',
        'success',
    )
    return redirect(url_for('orders.library'))

